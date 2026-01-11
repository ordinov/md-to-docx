#!/usr/bin/env python3
"""
Word (.docx) to Markdown converter - Self-contained version

Usage: python docx2md.py <file.docx>
Output: Creates <file.md> in the same directory

No manual installation required - handles everything automatically.
Works on Windows, Mac, and Linux.
"""

import sys
import os
import subprocess
import platform
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent.resolve()
VENV_DIR = SCRIPT_DIR / ".venv"


def get_python_executable():
    """Get the correct Python executable for the virtual environment."""
    if platform.system() == "Windows":
        return VENV_DIR / "Scripts" / "python.exe"
    else:
        return VENV_DIR / "bin" / "python"


def get_pip_executable():
    """Get the correct pip executable for the virtual environment."""
    if platform.system() == "Windows":
        return VENV_DIR / "Scripts" / "pip.exe"
    else:
        return VENV_DIR / "bin" / "pip"


def setup_environment():
    """Create virtual environment and install dependencies if needed."""
    python_exe = get_python_executable()
    pip_exe = get_pip_executable()

    # Check if venv exists and has python-docx
    if python_exe.exists():
        # Check if python-docx is installed
        result = subprocess.run(
            [str(python_exe), "-c", "import docx"],
            capture_output=True
        )
        if result.returncode == 0:
            return python_exe  # All good

    # Create venv if needed
    if not VENV_DIR.exists():
        print("Setting up environment (first run only)...")
        subprocess.run(
            [sys.executable, "-m", "venv", str(VENV_DIR)],
            check=True
        )

    # Install python-docx
    print("Installing dependencies...")
    subprocess.run(
        [str(pip_exe), "install", "-q", "python-docx"],
        check=True
    )
    print("Setup complete!\n")

    return python_exe


def run_in_venv():
    """Re-run this script inside the virtual environment."""
    python_exe = setup_environment()

    # Re-execute this script with the venv Python
    result = subprocess.run(
        [str(python_exe), __file__] + sys.argv[1:],
        env={**os.environ, "DOCX2MD_IN_VENV": "1"}
    )
    sys.exit(result.returncode)


# === Main conversion logic ===

def docx_to_md(docx_path):
    """Convert docx file to markdown, saving in same location."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    docx_path = Path(docx_path)

    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    if docx_path.suffix.lower() != '.docx':
        print(f"Error: File must have .docx extension: {docx_path}")
        print("Use md2docx.py to convert .md files to .docx")
        sys.exit(1)

    md_path = docx_path.with_suffix('.md')

    # Handle existing file
    if md_path.exists():
        response = input(f"File '{md_path.name}' already exists. Overwrite? (y/n): ").strip().lower()
        if response not in ('y', 's', 'yes', 'si'):
            # Find next available number
            counter = 1
            while True:
                new_path = md_path.parent / f"{md_path.stem} ({counter}){md_path.suffix}"
                if not new_path.exists():
                    md_path = new_path
                    break
                counter += 1

    doc = Document(docx_path)
    md_lines = []
    list_number_counter = 0  # Track numbered list position

    def extract_run_text(run):
        """Extract text from a run with formatting."""
        text = run.text
        if not text:
            return ""

        if run.bold and run.italic:
            return f"***{text}***"
        elif run.bold:
            return f"**{text}**"
        elif run.italic:
            return f"*{text}*"
        elif run.underline:
            # Underline often indicates links in converted docs
            return text
        return text

    def get_formatted_text(para):
        """Get paragraph text with inline formatting and hyperlinks."""
        result = []

        # Iterate through paragraph XML children to catch hyperlinks
        for child in para._p:
            if child.tag == qn('w:r'):
                # Regular run - find matching run object
                for run in para.runs:
                    if run._r is child:
                        result.append(extract_run_text(run))
                        break
            elif child.tag == qn('w:hyperlink'):
                # Hyperlink element
                r_id = child.get(qn('r:id'))
                url = None
                if r_id:
                    try:
                        rel = para.part.rels[r_id]
                        url = rel.target_ref
                    except (KeyError, AttributeError):
                        pass

                # Get text from runs inside hyperlink
                link_text = ""
                for run_elem in child.findall(qn('w:r')):
                    for text_elem in run_elem.findall(qn('w:t')):
                        if text_elem.text:
                            link_text += text_elem.text

                if url and link_text:
                    result.append(f"[{link_text}]({url})")
                elif link_text:
                    result.append(link_text)

        return "".join(result)

    def get_cell_formatted_text(cell, is_header=False):
        """Get cell text with inline formatting, joining paragraphs."""
        parts = []
        for para in cell.paragraphs:
            if is_header:
                # For header cells, get plain text (Word auto-bolds headers)
                para_text = para.text.strip()
            else:
                para_text = get_formatted_text(para)
            if para_text.strip():
                parts.append(para_text.strip())
        return " ".join(parts)

    def check_space_before(para):
        """Check if paragraph has space_before marker (indicating blank line in original)."""
        if para.paragraph_format.space_before:
            try:
                space_pt = para.paragraph_format.space_before.pt if hasattr(para.paragraph_format.space_before, 'pt') else 0
                if space_pt and space_pt >= 10:
                    return True
            except (AttributeError, TypeError):
                pass
        return False

    def process_paragraph(para):
        """Convert a paragraph to markdown."""
        nonlocal list_number_counter

        # Check heading style
        style_name = para.style.name if para.style else ""

        has_space = check_space_before(para)

        # Handle headings - return tuple (text, type)
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading', '').strip())
                text = para.text.strip()
                if level == 0 or style_name == 'Title':
                    return (f"# {text}", "heading_spaced" if has_space else True)
                else:
                    return (f"{'#' * (level + 1)} {text}", "heading_spaced" if has_space else True)
            except ValueError:
                pass

        if style_name == 'Title':
            return (f"# {para.text.strip()}", "heading_spaced" if has_space else True)

        # Handle lists
        if style_name == 'List Bullet':
            text = get_formatted_text(para)
            indent = ""
            if para.paragraph_format.left_indent:
                try:
                    indent_inches = para.paragraph_format.left_indent.inches if hasattr(para.paragraph_format.left_indent, 'inches') else 0
                    if indent_inches and indent_inches >= 0.4:
                        indent = "  "  # 2 spaces for sub-list
                        # Don't reset counter - this is a sub-list
                except (AttributeError, TypeError):
                    pass
            if not indent:
                list_number_counter = 0  # Only reset if not a sub-list
            return (f"{indent}- {text}", "spaced" if has_space else False)

        if style_name == 'List Number':
            list_number_counter += 1
            text = get_formatted_text(para)
            return (f"{list_number_counter}. {text}", "spaced" if has_space else False)

        # Check for horizontal rule (line of dashes or similar)
        text = para.text.strip()
        if text and all(c in '─-—' for c in text) and len(text) > 10:
            list_number_counter = 0  # Reset numbered list counter
            return ("---", "hr")  # Special marker for horizontal rule

        # Handle blockquotes (indented paragraphs)
        if para.paragraph_format.left_indent:
            try:
                indent_val = para.paragraph_format.left_indent.inches if hasattr(para.paragraph_format.left_indent, 'inches') else 0
                if indent_val and indent_val >= 0.4 and style_name not in ['List Bullet', 'List Number']:
                    text = get_formatted_text(para)
                    return (f"> {text}", "spaced" if has_space else False)
            except (AttributeError, TypeError):
                pass

        # Regular paragraph - reset list counter
        list_number_counter = 0
        formatted = get_formatted_text(para)
        return (formatted if formatted else "", "spaced" if has_space else False)

    def process_table(table):
        """Convert a table to markdown."""
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            is_header_row = (row_idx == 0)
            for cell in row.cells:
                # Get cell text - header row uses plain text (Word auto-bolds headers)
                cell_text = get_cell_formatted_text(cell, is_header=is_header_row)
                cells.append(cell_text)
            rows.append(cells)

        if not rows:
            return []

        def format_row(cells):
            """Format a row with proper spacing."""
            parts = []
            for cell in cells:
                if cell:
                    parts.append(f" {cell} ")
                else:
                    parts.append(" ")  # Empty cell = single space
            return "|" + "|".join(parts) + "|"

        md_table = []
        # Header row
        if rows:
            # Use header cell widths + 2 for separators (standard padding)
            col_widths = [max(len(cell) + 2, 3) for cell in rows[0]]

            md_table.append(format_row(rows[0]))
            separators = ["-" * w for w in col_widths]
            md_table.append("|" + "|".join(separators) + "|")
            # Data rows
            for row in rows[1:]:
                # Pad row if necessary
                while len(row) < len(rows[0]):
                    row.append("")
                md_table.append(format_row(row))

        return md_table

    # Process document elements in order
    prev_was_heading = False
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Find corresponding paragraph
            for para in doc.paragraphs:
                if para._element is element:
                    result = process_paragraph(para)
                    line, line_type = result

                    is_heading = line_type == True or line_type == "heading_spaced"
                    is_heading_spaced = line_type == "heading_spaced"
                    is_hr = line_type == "hr"
                    is_spaced = line_type == "spaced"

                    # Add blank line before horizontal rule
                    if is_hr and md_lines and md_lines[-1] != "":
                        md_lines.append("")

                    # Add blank line before paragraph/heading that had space_before
                    if (is_spaced or is_heading_spaced) and md_lines and md_lines[-1] != "":
                        md_lines.append("")

                    # Add blank line after previous heading/hr if this is content
                    if prev_was_heading and line and not is_hr and not is_heading:
                        if md_lines and md_lines[-1] != "":
                            md_lines.append("")

                    if line or (md_lines and md_lines[-1] != ""):
                        md_lines.append(line)

                    prev_was_heading = is_heading or is_hr
                    break
        elif element.tag.endswith('tbl'):
            # Find corresponding table
            for table in doc.tables:
                if table._tbl is element:
                    # Add blank line before table if needed
                    if md_lines and md_lines[-1] != "":
                        md_lines.append("")
                    md_lines.extend(process_table(table))
                    md_lines.append("")
                    prev_was_heading = False
                    list_number_counter = 0  # Reset after table
                    break

    # Clean up multiple blank lines
    cleaned = []
    prev_blank = False
    for line in md_lines:
        is_blank = line == ""
        if is_blank and prev_blank:
            continue
        cleaned.append(line)
        prev_blank = is_blank

    # Remove trailing blank lines
    while cleaned and cleaned[-1] == "":
        cleaned.pop()

    content = "\n".join(cleaned) + "\n"  # Add trailing newline

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)

    print(f"Converted: {docx_path.name} -> {md_path.name}")
    return md_path


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python docx2md.py <file.docx>")
        print("       Creates <file.md> in the same directory")
        sys.exit(1)

    # If not in venv, set it up and re-run
    if not os.environ.get("DOCX2MD_IN_VENV"):
        run_in_venv()
    else:
        docx_to_md(sys.argv[1])
