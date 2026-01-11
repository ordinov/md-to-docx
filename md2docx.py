#!/usr/bin/env python3
"""
Markdown to Word (.docx) converter - Self-contained version

Usage: python md2docx.py <file.md>
Output: Creates <file.docx> in the same directory

No manual installation required - handles everything automatically.
Works on Windows, Mac, and Linux.
"""

import sys
import os
import subprocess
import platform
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent.resolve()
VENV_DIR = SCRIPT_DIR / ".venv"


def get_python_executable():
    """Get the correct Python executable for the virtual environment."""
    if platform.system() == "Windows":
        return VENV_DIR / "Scripts" / "python.exe"
    else:
        return VENV_DIR / "bin" / "python"


def get_pip_executable():
    """Get the correct pip executable for the virtual environment."""
    if platform.system() == "Windows":
        return VENV_DIR / "Scripts" / "pip.exe"
    else:
        return VENV_DIR / "bin" / "pip"


def setup_environment():
    """Create virtual environment and install dependencies if needed."""
    python_exe = get_python_executable()
    pip_exe = get_pip_executable()

    # Check if venv exists and has python-docx
    if python_exe.exists():
        # Check if python-docx is installed
        result = subprocess.run(
            [str(python_exe), "-c", "import docx"],
            capture_output=True
        )
        if result.returncode == 0:
            return python_exe  # All good

    # Create venv if needed
    if not VENV_DIR.exists():
        print("Setting up environment (first run only)...")
        subprocess.run(
            [sys.executable, "-m", "venv", str(VENV_DIR)],
            check=True
        )

    # Install python-docx
    print("Installing dependencies...")
    subprocess.run(
        [str(pip_exe), "install", "-q", "python-docx"],
        check=True
    )
    print("Setup complete!\n")

    return python_exe


def run_in_venv():
    """Re-run this script inside the virtual environment."""
    python_exe = setup_environment()

    # Re-execute this script with the venv Python
    result = subprocess.run(
        [str(python_exe), __file__] + sys.argv[1:],
        env={**os.environ, "MD2DOCX_IN_VENV": "1"}
    )
    sys.exit(result.returncode)


# === Main conversion logic ===

def md_to_docx(md_path):
    """Convert markdown file to docx, saving in same location."""
    import re
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    md_path = Path(md_path)

    if not md_path.exists():
        print(f"Error: File not found: {md_path}")
        sys.exit(1)

    if md_path.suffix.lower() != '.md':
        print(f"Error: File must have .md extension: {md_path}")
        print("Use docx2md.py to convert .docx files to .md")
        sys.exit(1)

    docx_path = md_path.with_suffix('.docx')

    # Handle existing file
    if docx_path.exists():
        response = input(f"File '{docx_path.name}' already exists. Overwrite? (y/n): ").strip().lower()
        if response not in ('y', 's', 'yes', 'si'):
            # Find next available number
            counter = 1
            while True:
                new_path = docx_path.parent / f"{docx_path.stem} ({counter}){docx_path.suffix}"
                if not new_path.exists():
                    docx_path = new_path
                    break
                counter += 1

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []

    def create_table(doc, data):
        if not data:
            return
        rows = len(data)
        cols = max(len(row) for row in data)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for idx, row_data in enumerate(data):
            row = table.rows[idx]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_formatted_text(p, cell_text)
                    if idx == 0:
                        for run in p.runs:
                            run.bold = True
        doc.add_paragraph()

    def add_hyperlink(paragraph, url, text):
        """Add a hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Blue color for hyperlink
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)

        # Underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)

        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def add_formatted_text(paragraph, text):
        pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        parts = re.split(pattern, text)
        idx = 0
        while idx < len(parts):
            if idx + 2 < len(parts) and idx % 3 == 0:
                if parts[idx]:
                    add_styled_text(paragraph, parts[idx])
                # Add hyperlink with URL
                link_text = parts[idx + 1]
                link_url = parts[idx + 2]
                add_hyperlink(paragraph, link_url, link_text)
                idx += 3
            else:
                if parts[idx]:
                    add_styled_text(paragraph, parts[idx])
                idx += 1

    def add_styled_text(paragraph, text):
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                italic_parts = re.split(r'(\*[^*]+\*)', part)
                for ip in italic_parts:
                    if ip.startswith('*') and ip.endswith('*') and len(ip) > 2:
                        run = paragraph.add_run(ip[1:-1])
                        run.italic = True
                    else:
                        if ip:
                            paragraph.add_run(ip)

    prev_blank = False  # Track if previous line was blank
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            if in_table and table_data:
                create_table(doc, table_data)
                table_data = []
                in_table = False
            prev_blank = True
            i += 1
            continue

        if line.strip() == '---':
            if not in_table:
                doc.add_paragraph('─' * 50)
            prev_blank = False
            i += 1
            continue

        if line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=0)
            if prev_blank:
                h.paragraph_format.space_before = Pt(12)
            prev_blank = False
            i += 1
            continue
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=1)
            if prev_blank:
                h.paragraph_format.space_before = Pt(12)
            prev_blank = False
            i += 1
            continue
        elif line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=2)
            if prev_blank:
                h.paragraph_format.space_before = Pt(12)
            prev_blank = False
            i += 1
            continue
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:].strip(), level=3)
            if prev_blank:
                h.paragraph_format.space_before = Pt(12)
            prev_blank = False
            i += 1
            continue

        if line.strip().startswith('|') and line.strip().endswith('|'):
            in_table = True
            prev_blank = False
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            table_data.append(cells)
            i += 1
            continue
        elif in_table:
            create_table(doc, table_data)
            table_data = []
            in_table = False

        # Sub-list: 2+ spaces before "- " (must check BEFORE regular bullet)
        if re.match(r'^  +- ', line):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, text)
            prev_blank = False
            i += 1
            continue

        if line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            if prev_blank:
                p.paragraph_format.space_before = Pt(12)
            add_formatted_text(p, text)
            prev_blank = False
            i += 1
            continue

        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            # Mark with space_before if preceded by blank line (for round-trip)
            if prev_blank:
                p.paragraph_format.space_before = Pt(12)
            add_formatted_text(p, text)
            prev_blank = False
            i += 1
            continue

        if line.strip().startswith('> '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            if prev_blank:
                p.paragraph_format.space_before = Pt(12)
            add_formatted_text(p, text)
            prev_blank = False
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            if prev_blank:
                p.paragraph_format.space_before = Pt(12)
            add_formatted_text(p, line.strip())

        prev_blank = False
        i += 1

    if table_data:
        create_table(doc, table_data)

    doc.save(docx_path)
    print(f"Converted: {md_path.name} -> {docx_path.name}")
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python md2docx.py <file.md>")
        print("       Creates <file.docx> in the same directory")
        sys.exit(1)

    # If not in venv, set it up and re-run
    if not os.environ.get("MD2DOCX_IN_VENV"):
        run_in_venv()
    else:
        md_to_docx(sys.argv[1])
